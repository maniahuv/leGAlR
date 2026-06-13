from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from langchain_core.documents import Document

REPO_ROOT = Path(__file__).resolve().parents[2]


REQUIRED_MANIFEST_FIELDS = ("filename", "doc_id", "title")
SUPPORTED_SOURCE_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def resolve_repo_path(path: str | Path | None, default: str | None = None) -> Path:
    """Resolve a config path relative to the repository root."""
    raw = Path(path or default or ".")
    if raw.is_absolute():
        return raw
    return REPO_ROOT / raw


def read_jsonl(path: str | Path) -> list[dict[str, Any]]:
    file_path = resolve_repo_path(path)
    if not file_path.exists():
        return []

    rows: list[dict[str, Any]] = []
    with file_path.open("r", encoding="utf-8") as f:
        for line_no, line in enumerate(f, start=1):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            try:
                obj = json.loads(line)
            except json.JSONDecodeError as exc:
                raise ValueError(f"Invalid JSONL at {file_path}:{line_no}: {exc}") from exc
            if not isinstance(obj, dict):
                raise ValueError(f"Invalid JSONL at {file_path}:{line_no}: each line must be an object")
            rows.append(obj)
    return rows


def write_jsonl(path: str | Path, rows: Iterable[dict[str, Any]]) -> None:
    file_path = resolve_repo_path(path)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    with file_path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False, sort_keys=False) + "\n")


def write_json(path: str | Path, payload: dict[str, Any]) -> None:
    file_path = resolve_repo_path(path)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_doc_id(value: str) -> str:
    value = unicodedata.normalize("NFC", value or "").strip()
    value = value.replace("/", "_").replace("-", "_")
    value = re.sub(r"[^0-9A-Za-z_]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value


def _metadata_value(row: dict[str, Any], key: str, default: str = "") -> str:
    value = row.get(key, default)
    if value is None:
        return ""
    return str(value).strip()


def _as_bool(value: Any, default: bool = True) -> bool:
    if value is None or value == "":
        return default
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "y", "có", "co", "còn hiệu lực", "con hieu luc"}


def sanitize_metadata(metadata: dict[str, Any]) -> dict[str, str | int | float | bool]:
    """Chroma only accepts simple scalar metadata values."""
    sanitized: dict[str, str | int | float | bool] = {}
    for key, value in (metadata or {}).items():
        if value is None:
            sanitized[key] = ""
        elif isinstance(value, (str, int, float, bool)):
            sanitized[key] = value
        else:
            sanitized[key] = json.dumps(value, ensure_ascii=False)
    return sanitized


def validate_manifest_rows(rows: list[dict[str, Any]], manifest_path: str | Path) -> None:
    if not rows:
        raise ValueError(
            f"Manifest is empty or missing: {resolve_repo_path(manifest_path)}. "
            "Add one JSON object per source file. See data/raw/family_law/manifest.example.jsonl."
        )

    seen_doc_ids: set[str] = set()
    for i, row in enumerate(rows, start=1):
        missing = [field for field in REQUIRED_MANIFEST_FIELDS if not _metadata_value(row, field)]
        if missing:
            raise ValueError(f"Manifest row {i} is missing required fields: {missing}")
        doc_id = _metadata_value(row, "doc_id")
        if doc_id in seen_doc_ids:
            raise ValueError(f"Duplicate doc_id in manifest: {doc_id}")
        seen_doc_ids.add(doc_id)

        filename = _metadata_value(row, "filename")
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_SOURCE_EXTENSIONS:
            raise ValueError(
                f"Manifest row {i} has unsupported file extension '{ext}' for file {filename}. "
                f"Supported: {sorted(SUPPORTED_SOURCE_EXTENSIONS)}"
            )


def _validate_source_file(file_path: Path) -> dict[str, Any]:
    if not file_path.exists():
        raise FileNotFoundError(f"Source file not found: {file_path}")
    if not file_path.is_file():
        raise FileNotFoundError(f"Source path is not a file: {file_path}")

    size_bytes = file_path.stat().st_size
    if size_bytes <= 0:
        raise ValueError(f"Source file is empty: {file_path}")

    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_SOURCE_EXTENSIONS:
        raise ValueError(f"Unsupported source file type: {file_path.name}. Supported: {sorted(SUPPORTED_SOURCE_EXTENSIONS)}")

    header = file_path.read_bytes()[:2048]
    header_lower = header.lower()
    if b"<html" in header_lower or b"<!doctype html" in header_lower:
        raise ValueError(
            f"{file_path.name} looks like an HTML page saved with extension {ext}. "
            "Download the actual attachment file from vbpl.vn/vanban.chinhphu.vn again."
        )

    if ext == ".pdf" and b"%pdf" not in header_lower[:1024]:
        raise ValueError(
            f"{file_path.name} does not look like a valid PDF file because it has no %PDF header. "
            "The download may be incomplete/corrupted; delete it and download again."
        )

    if ext == ".docx" and not header.startswith(b"PK"):
        raise ValueError(
            f"{file_path.name} does not look like a valid DOCX file. "
            "If it is an old .doc file, keep the .doc extension or save it as .docx from Word."
        )

    return {"source_path": str(file_path), "file_size_bytes": size_bytes, "file_extension": ext}


def _extract_pdf_text_with_pypdf(file_path: Path) -> tuple[str, dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path), strict=False)
    pages: list[str] = []
    empty_pages = 0
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if not text.strip():
            empty_pages += 1
        pages.append(f"\n\n[PAGE {page_index}]\n{text.strip()}")

    content = "\n".join(pages).strip()
    report = {
        "extractor": "pypdf",
        "page_count": len(reader.pages),
        "empty_pages": empty_pages,
        "char_count": len(content),
    }
    return content, report


def _extract_pdf_text_with_pymupdf(file_path: Path) -> tuple[str, dict[str, Any]]:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:  # pragma: no cover
        raise ImportError("Missing fallback dependency PyMuPDF. Install with: pip install PyMuPDF") from exc

    pages: list[str] = []
    empty_pages = 0
    with fitz.open(str(file_path)) as doc:
        page_count = doc.page_count
        for page_index, page in enumerate(doc, start=1):
            try:
                text = page.get_text("text") or ""
            except Exception:
                text = ""
            if not text.strip():
                empty_pages += 1
            pages.append(f"\n\n[PAGE {page_index}]\n{text.strip()}")

    content = "\n".join(pages).strip()
    report = {
        "extractor": "pymupdf",
        "page_count": page_count,
        "empty_pages": empty_pages,
        "char_count": len(content),
    }
    return content, report


def extract_pdf_text(pdf_path: str | Path) -> tuple[str, dict[str, Any]]:
    """Extract text from PDF with pypdf first, then PyMuPDF fallback."""
    file_path = resolve_repo_path(pdf_path)
    base_report = _validate_source_file(file_path)

    if file_path.suffix.lower() != ".pdf":
        raise ValueError(f"extract_pdf_text only accepts .pdf files, got: {file_path.name}")

    pypdf_error = ""
    try:
        content, report = _extract_pdf_text_with_pypdf(file_path)
        report.update(base_report)
        return content, report
    except Exception as exc:
        pypdf_error = f"{type(exc).__name__}: {exc}"

    try:
        content, report = _extract_pdf_text_with_pymupdf(file_path)
        report.update(base_report)
        report["pypdf_error"] = pypdf_error
        return content, report
    except Exception as exc:
        pymupdf_error = f"{type(exc).__name__}: {exc}"
        raise ValueError(
            "Could not extract text from PDF.\n"
            f"File: {file_path}\n"
            f"pypdf error: {pypdf_error}\n"
            f"PyMuPDF error: {pymupdf_error}\n"
            "Fix: delete this PDF and download it again. If the PDF is scanned, use the official Word/DOCX file "
            "or OCR before ingest."
        ) from exc


def _extract_docx_text(file_path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:  # pragma: no cover
        raise ImportError("Missing dependency python-docx. Install with: pip install python-docx") from exc

    doc = DocxDocument(str(file_path))
    blocks: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    table_count = 0
    for table in doc.tables:
        table_count += 1
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append("\n".join(rows))

    content = "\n\n".join(blocks).strip()
    report = {
        "extractor": "python-docx",
        "paragraph_count": len(doc.paragraphs),
        "table_count": table_count,
        "char_count": len(content),
    }
    return content, report


def _extract_doc_text_with_soffice(file_path: Path) -> tuple[str, dict[str, Any]]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice is not available in PATH")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        cmd = [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", str(tmp_path), str(file_path)]
        completed = subprocess.run(cmd, capture_output=True, text=True, timeout=90)
        if completed.returncode != 0:
            raise RuntimeError(f"soffice failed: {completed.stderr or completed.stdout}")
        txt_files = list(tmp_path.glob("*.txt"))
        if not txt_files:
            raise RuntimeError("soffice did not produce a .txt output file")
        raw = txt_files[0].read_bytes()

    for enc in ("utf-8", "utf-16", "cp1258", "cp1252"):
        try:
            content = raw.decode(enc).strip()
            break
        except UnicodeDecodeError:
            continue
    else:
        content = raw.decode("utf-8", errors="ignore").strip()

    return content, {"extractor": "soffice", "char_count": len(content)}


def _extract_doc_text_with_word(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Fallback for old .doc files on Windows with Microsoft Word + pywin32 installed."""
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError("pywin32/Microsoft Word fallback is not available") from exc

    with tempfile.TemporaryDirectory() as tmpdir:
        out_path = Path(tmpdir) / f"{file_path.stem}.txt"
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(file_path.resolve()))
            # 7 = wdFormatUnicodeText in many Word versions; robust enough for Vietnamese.
            doc.SaveAs2(str(out_path), FileFormat=7)
            doc.Close(False)
        finally:
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()

        raw = out_path.read_bytes()

    for enc in ("utf-16", "utf-8", "cp1258", "cp1252"):
        try:
            content = raw.decode(enc).strip()
            break
        except UnicodeDecodeError:
            continue
    else:
        content = raw.decode("utf-8", errors="ignore").strip()

    return content, {"extractor": "microsoft_word_com", "char_count": len(content)}


def _extract_doc_text(file_path: Path) -> tuple[str, dict[str, Any]]:
    errors: list[str] = []
    try:
        return _extract_doc_text_with_soffice(file_path)
    except Exception as exc:
        errors.append(f"soffice: {type(exc).__name__}: {exc}")
    try:
        return _extract_doc_text_with_word(file_path)
    except Exception as exc:
        errors.append(f"word-com: {type(exc).__name__}: {exc}")
    raise ValueError(
        f"Cannot extract old .doc file: {file_path}\n"
        + "\n".join(errors)
        + "\nFix: open this file in Microsoft Word and Save As .docx, then set filename to the .docx file in manifest.jsonl."
    )


def _extract_plain_text(file_path: Path) -> tuple[str, dict[str, Any]]:
    raw = file_path.read_bytes()
    for enc in ("utf-8", "utf-16", "cp1258", "cp1252"):
        try:
            content = raw.decode(enc).strip()
            break
        except UnicodeDecodeError:
            continue
    else:
        content = raw.decode("utf-8", errors="ignore").strip()
    return content, {"extractor": "plain_text", "char_count": len(content)}


def extract_source_text(source_path: str | Path) -> tuple[str, dict[str, Any]]:
    """Extract text from official local sources: PDF, DOCX, DOC, TXT, MD."""
    file_path = resolve_repo_path(source_path)
    base_report = _validate_source_file(file_path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        content, report = extract_pdf_text(file_path)
    elif ext == ".docx":
        content, report = _extract_docx_text(file_path)
    elif ext == ".doc":
        content, report = _extract_doc_text(file_path)
    elif ext in {".txt", ".md"}:
        content, report = _extract_plain_text(file_path)
    else:  # validate_source_file should catch this.
        raise ValueError(f"Unsupported source file type: {file_path.name}")

    report.update(base_report)
    report["source_file"] = file_path.name
    return content, report


def build_metadata_record(row: dict[str, Any], source_path: Path) -> dict[str, Any]:
    doc_id = _metadata_value(row, "doc_id") or normalize_doc_id(Path(_metadata_value(row, "filename")).stem)
    tinh_trang = _metadata_value(row, "tinh_trang_hieu_luc", "Còn hiệu lực")
    is_current = _as_bool(row.get("is_current"), default=("hết hiệu lực" not in tinh_trang.lower()))

    source_type = _metadata_value(row, "source_type")
    if not source_type:
        ext = source_path.suffix.lower().lstrip(".")
        source_type = f"official_{ext}"

    return {
        "doc_id": doc_id,
        "id": doc_id,  # compatibility with HF-style code
        "title": _metadata_value(row, "title"),
        "so_ky_hieu": _metadata_value(row, "so_ky_hieu"),
        "loai_van_ban": _metadata_value(row, "loai_van_ban"),
        "ngay_ban_hanh": _metadata_value(row, "ngay_ban_hanh"),
        "ngay_co_hieu_luc": _metadata_value(row, "ngay_co_hieu_luc"),
        "ngay_het_hieu_luc": _metadata_value(row, "ngay_het_hieu_luc"),
        "tinh_trang_hieu_luc": tinh_trang,
        "is_current": is_current,
        "nganh": _metadata_value(row, "nganh"),
        "linh_vuc": _metadata_value(row, "linh_vuc", "Hôn nhân và gia đình"),
        "co_quan_ban_hanh": _metadata_value(row, "co_quan_ban_hanh"),
        "chuc_danh": _metadata_value(row, "chuc_danh"),
        "nguoi_ky": _metadata_value(row, "nguoi_ky"),
        "pham_vi": _metadata_value(row, "pham_vi"),
        "thong_tin_ap_dung": _metadata_value(row, "thong_tin_ap_dung"),
        "corpus_role": _metadata_value(row, "corpus_role", "core"),
        "source": "local_official_file",
        "source_type": source_type,
        "source_url": _metadata_value(row, "source_url"),
        "filename": _metadata_value(row, "filename"),
        "local_file_path": str(source_path),
        # Keep this old key for backward compatibility with older scripts.
        "local_pdf_path": str(source_path) if source_path.suffix.lower() == ".pdf" else "",
    }


def build_document_header(metadata: dict[str, Any]) -> str:
    lines = [
        f"Văn bản: {metadata.get('title', '')}",
        f"Số hiệu: {metadata.get('so_ky_hieu', '')}",
        f"Loại văn bản: {metadata.get('loai_van_ban', '')}",
        f"Cơ quan ban hành: {metadata.get('co_quan_ban_hanh', '')}",
        f"Lĩnh vực: {metadata.get('linh_vuc', '')}",
        f"Vai trò corpus: {metadata.get('corpus_role', '')}",
        f"Ngày ban hành: {metadata.get('ngay_ban_hanh', '')}",
        f"Ngày có hiệu lực: {metadata.get('ngay_co_hieu_luc', '')}",
        f"Ngày hết hiệu lực: {metadata.get('ngay_het_hieu_luc', '')}",
        f"Tình trạng hiệu lực: {metadata.get('tinh_trang_hieu_luc', '')}",
        f"Nguồn: {metadata.get('source_url', '')}",
        f"File nguồn: {metadata.get('filename', '')}",
        "Nội dung:",
    ]
    return "\n".join(lines).strip() + "\n"


@dataclass
class LocalPdfCorpus:
    # Historical class name kept to avoid breaking imports. It now stores PDF/DOCX/DOC/TXT sources.
    documents: list[Document]
    metadata_records: list[dict[str, Any]]
    content_records: list[dict[str, Any]]
    extraction_reports: list[dict[str, Any]]


def _resolve_manifest_file_path(raw_dir: Path, filename: str) -> Path:
    raw = Path(filename)
    if raw.is_absolute():
        return raw
    # If the manifest uses subfolders like pdfs/file.pdf or word/file.docx, respect that relative path.
    direct = resolve_repo_path(raw)
    if direct.exists():
        return direct
    return raw_dir / raw


def load_local_family_law_corpus(config) -> LocalPdfCorpus:
    dataset_cfg = getattr(config, "dataset", config)
    raw_dir = resolve_repo_path(getattr(dataset_cfg, "raw_dir", "data/raw/family_law/pdfs"))
    manifest_path = resolve_repo_path(getattr(dataset_cfg, "manifest_path", "data/raw/family_law/manifest.jsonl"))

    manifest_rows = read_jsonl(manifest_path)
    validate_manifest_rows(manifest_rows, manifest_path)

    documents: list[Document] = []
    metadata_records: list[dict[str, Any]] = []
    content_records: list[dict[str, Any]] = []
    extraction_reports: list[dict[str, Any]] = []

    for row in manifest_rows:
        filename = _metadata_value(row, "filename")
        source_path = _resolve_manifest_file_path(raw_dir, filename)
        metadata = build_metadata_record(row, source_path)
        try:
            content_text, report = extract_source_text(source_path)
        except Exception as exc:
            raise RuntimeError(
                f"Failed to extract text for doc_id={metadata.get('doc_id')} "
                f"title={metadata.get('title')} file={source_path}: {exc}"
            ) from exc

        min_chars = int(getattr(dataset_cfg, "min_pdf_chars", getattr(dataset_cfg, "min_source_chars", 100)))
        if len(content_text.strip()) < min_chars:
            hint = ""
            if source_path.suffix.lower() == ".pdf":
                hint = " This PDF may be scanned image; use the official Word/DOCX file or OCR."
            raise ValueError(f"Extracted text is too short for {source_path}.{hint}")

        doc_id = str(metadata["doc_id"])
        content_record = {
            "doc_id": doc_id,
            "id": doc_id,
            "content_text": content_text,
            "content_format": f"{source_path.suffix.lower().lstrip('.')}_extracted_text",
            "source_file": filename,
            "source_pdf": filename if source_path.suffix.lower() == ".pdf" else "",
            "source_url": metadata.get("source_url", ""),
        }
        report.update({"doc_id": doc_id, "title": metadata.get("title", "")})

        metadata_records.append(metadata)
        content_records.append(content_record)
        extraction_reports.append(report)

        page_content = build_document_header(metadata) + content_text
        documents.append(Document(page_content=page_content, metadata=sanitize_metadata(metadata)))

    return LocalPdfCorpus(
        documents=documents,
        metadata_records=metadata_records,
        content_records=content_records,
        extraction_reports=extraction_reports,
    )


def load_local_family_law_documents(config) -> list[Document]:
    return load_local_family_law_corpus(config).documents
